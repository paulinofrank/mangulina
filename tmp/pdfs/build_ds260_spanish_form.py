from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Mangulina\output\Formulario_DS-260_en_espanol_Josue.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT = "E8EEF5"
GRAY = "666666"
LINE = "9AA7B5"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(.72)
sec.bottom_margin = Inches(.68)
sec.left_margin = Inches(.78)
sec.right_margin = Inches(.78)
sec.header_distance = Inches(.32)
sec.footer_distance = Inches(.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ("Title", 24, NAVY, 0, 5),
    ("Subtitle", 11, GRAY, 0, 14),
    ("Heading 1", 15, NAVY, 14, 7),
    ("Heading 2", 12, BLUE, 10, 4),
    ("Heading 3", 10.5, NAVY, 7, 3),
]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def bottom_border(p, color=LINE, size="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)

def add_field(label=None, lines=1):
    for i in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.keep_together = True
        if label and i == 0:
            r = p.add_run(label + "  ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(" ")
        bottom_border(p)

def add_q(prompt, yesno=False, lines=1, hint=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(prompt)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("222222")
    if hint:
        hr = p.add_run("\n" + hint)
        hr.italic = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor.from_string(GRAY)
    if yesno:
        a = doc.add_paragraph("Sí [  ]     No [  ]")
        a.paragraph_format.left_indent = Inches(.18)
        a.paragraph_format.space_after = Pt(5)
        a.paragraph_format.keep_together = True
    else:
        add_field(lines=lines)

def section(title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.left_indent = Inches(-.08)
    p.paragraph_format.right_indent = Inches(-.08)
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(7)
    shade_paragraph(p, LIGHT)
    r = p.add_run("  " + title)
    r.font.color.rgb = RGBColor.from_string(NAVY)

def sub(title):
    doc.add_paragraph(title, style="Heading 2")

def page_break():
    doc.add_page_break()

# Header / footer
hp = sec.header.paragraphs[0]
hp.text = "GUÍA DE PREPARACIÓN - FORMULARIO DS-260"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs:
    r.font.name = "Arial"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)
bottom_border(hp, "D9E0E8", "4")

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Página ")
r.font.name = "Arial"; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# Opening block
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(9)
p.paragraph_format.space_after = Pt(1)
r = p.add_run("FORMULARIO DE PREPARACIÓN")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(style="Title")
p.add_run("Solicitud DS-260").bold = True
p = doc.add_paragraph(style="Subtitle")
p.add_run("Preguntas en español para completar manualmente antes de ingresar la información en línea")

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(10)
shade_paragraph(note, "F4F6F9")
r = note.add_run("Importante: ")
r.bold = True; r.font.color.rgb = RGBColor.from_string(NAVY)
note.add_run("Esta es una guía de preparación basada en el archivo adjunto; no sustituye el formulario oficial. Responda con información completa y veraz. Si marca “Sí” en una pregunta que requiere explicación, use el espacio adicional correspondiente.")

add_field("Nombre del solicitante")
add_field("Número de caso")
add_field("Número de confirmación")
add_field("Fecha en que se completó esta guía")

section("1. Información personal, dirección y contacto")
sub("Datos personales")
for label in [
    "Apellidos tal como aparecen en el pasaporte", "Nombres", "Nombre completo en su idioma o alfabeto nativo",
]: add_field(label)
add_q("¿Ha utilizado otros nombres (incluidos apodos, nombres de soltero/a, religiosos, profesionales o alias)?", True)
add_field("Si respondió Sí, indique todos los demás nombres", 2)
add_q("Sexo", False)
p = doc.paragraphs[-1]; p.text = "Masculino [  ]     Femenino [  ]"
add_field("Estado civil actual")
add_field("Fecha de nacimiento (día/mes/año)")
add_field("Ciudad de nacimiento")
add_field("Estado/provincia de nacimiento")
add_field("País/región de nacimiento")
add_field("País/región de origen o nacionalidad")

sub("Documento de viaje")
add_field("Tipo de documento (por ejemplo, pasaporte)")
add_field("Número o identificación del documento")
add_field("País/autoridad que lo expidió")
add_field("Fecha de expedición")
add_field("Fecha de vencimiento")
add_q("¿Tiene o ha tenido alguna nacionalidad distinta de la indicada anteriormente?", True)
add_field("Si respondió Sí, indique el país y los detalles", 2)

sub("Dirección actual")
for label in ["Calle y número", "Apartamento/unidad", "Ciudad", "Estado/provincia", "Código postal", "País/región", "¿Desde qué fecha vive allí?"]:
    add_field(label)
add_q("¿Ha vivido en algún otro lugar desde que cumplió 16 años?", True)
add_field("Si respondió Sí, enumere cada dirección y las fechas desde/hasta", 4)

sub("Teléfonos, correo electrónico y redes sociales")
for label in ["Teléfono principal", "Teléfono secundario", "Teléfono del trabajo"]: add_field(label)
add_q("¿Ha utilizado otros números de teléfono?", True)
add_field("Si respondió Sí, enumérelos", 2)
add_field("Correo electrónico actual")
add_q("¿Ha utilizado otras direcciones de correo electrónico?", True)
add_field("Si respondió Sí, enumérelas", 2)
add_field("Proveedor o plataforma de red social")
add_field("Nombre de usuario o identificador en esa red social")
add_q("¿Ha utilizado otras redes sociales?", True)
add_field("Si respondió Sí, indique plataforma e identificador de cada una", 3)

sub("Dirección postal y domicilio permanente en Estados Unidos")
add_q("¿Su dirección postal es la misma que su dirección actual?", True)
add_field("Si respondió No, indique la dirección postal completa", 3)
add_field("Nombre de la persona que vive actualmente en el domicilio permanente de EE. UU.")
for label in ["Calle y número en EE. UU.", "Apartamento/unidad", "Ciudad", "Estado", "Código postal", "Teléfono"]: add_field(label)
add_q("¿Desea que su Tarjeta de Residente Permanente (Green Card) sea enviada a esta dirección?", True)
add_field("Si respondió No, indique la dirección correcta para recibirla", 3)

section("2. Información familiar")
sub("Padre")
for label in ["Apellidos del padre", "Nombres del padre", "Fecha de nacimiento", "Ciudad de nacimiento", "Estado/provincia de nacimiento", "País/región de nacimiento"]: add_field(label)
add_q("¿Su padre vive actualmente?", True)
add_field("Dirección actual del padre (calle, ciudad, estado/provincia, código postal y país)", 3)

sub("Madre")
for label in ["Apellidos de nacimiento de la madre", "Nombres de la madre", "Fecha de nacimiento", "Ciudad de nacimiento", "Estado/provincia de nacimiento", "País/región de nacimiento"]: add_field(label)
add_q("¿Su madre vive actualmente?", True)
add_q("¿La dirección de su madre es la misma que la de su padre?", True)
add_field("Si no es la misma, indique la dirección actual completa de su madre", 3)

sub("Cónyuges e hijos")
add_q("¿Ha tenido algún cónyuge anterior?", True)
add_field("Si respondió Sí, indique nombre completo, fecha y lugar de nacimiento, nacionalidad, fechas del matrimonio y cómo terminó", 4)
add_q("¿Tiene hijos?", True)
add_field("Si respondió Sí, indique por cada hijo: nombre completo, fecha y lugar de nacimiento, sexo, nacionalidad y dirección", 5)

section("3. Viajes anteriores a Estados Unidos")
add_q("¿Ha estado alguna vez en Estados Unidos?", True)
add_field("Si respondió Sí, indique fechas de llegada, duración y estatus de cada viaje", 4)
add_q("¿Alguna vez le han expedido una visa de Estados Unidos?", True)
add_field("Si respondió Sí, indique tipo de visa, número, fecha y lugar de expedición", 3)
add_q("¿Alguna vez le han negado una visa de Estados Unidos, le han negado la admisión al país o retiró su solicitud de admisión en un puerto de entrada?", True)
add_field("Si respondió Sí, explique cuándo, dónde y qué ocurrió", 4)

section("4. Trabajo, educación y capacitación")
sub("Ocupación y empleo actual")
add_field("Ocupación principal")
add_field("Si seleccionó “Otra”, especifique")
add_field("Nombre del empleador o institución educativa actual")
for label in ["Dirección", "Ciudad", "Estado/provincia", "Código postal", "País/región"]: add_field(label)
add_q("¿Tiene otras ocupaciones?", True)
add_field("Si respondió Sí, especifique")
add_field("¿En qué ocupación piensa trabajar en Estados Unidos?")
add_field("Si seleccionó “Otra”, especifique")
add_q("¿Ha tenido empleos anteriores?", True)
add_field("Si respondió Sí, indique por cada empleo: empleador, dirección, cargo, funciones y fechas desde/hasta", 5)

sub("Educación")
add_q("¿Ha asistido a alguna institución educativa?", True)
add_field("Número de instituciones a las que asistió")
for n in range(1, 4):
    doc.add_paragraph(f"Institución educativa {n}", style="Heading 3")
    for label in ["Nombre", "Dirección", "Ciudad", "Estado/provincia", "Código postal", "País/región", "Área o curso de estudio", "Título o diploma", "Fecha de asistencia desde", "Fecha de asistencia hasta"]:
        add_field(label)

sub("Viajes, servicio, organizaciones, habilidades e idiomas")
add_q("¿Ha viajado a algún país o región durante los últimos cinco años?", True)
add_field("Si respondió Sí, enumere los países o regiones visitados", 3)
add_q("¿Ha prestado servicio militar alguna vez?", True)
add_field("Si respondió Sí, indique país, rama, rango, especialidad y fechas", 3)
add_q("¿Ha pertenecido, contribuido o trabajado para alguna organización profesional, social o benéfica?", True)
add_field("Si respondió Sí, indique nombre, ubicación, función y fechas", 3)
add_q("¿Tiene habilidades o capacitación especializada, incluida experiencia con armas de fuego, explosivos o materiales nucleares, biológicos o químicos?", True)
add_field("Si respondió Sí, explique", 3)
add_q("¿Ha servido, pertenecido o estado involucrado con una unidad paramilitar, grupo de vigilantes, grupo rebelde, guerrilla u organización insurgente?", True)
add_field("Si respondió Sí, explique", 3)
add_q("¿Puede hablar o leer idiomas distintos de su idioma nativo?", True)
add_field("Si respondió Sí, enumere los idiomas e indique si habla y/o lee cada uno", 3)

section("5. Información del peticionario")
add_field("El peticionario es mi (relación con usted)")
add_field("Apellidos del peticionario")
add_field("Nombres del peticionario")
for label in ["Dirección", "Apartamento/unidad", "Ciudad", "Estado/provincia", "Código postal", "País/región", "Teléfono", "Teléfono móvil/celular", "Correo electrónico"]: add_field(label)

section("6. Seguridad y antecedentes")
p = doc.add_paragraph("Marque Sí o No en cada pregunta. Para toda respuesta afirmativa, incluya una explicación completa en el espacio indicado y reúna los documentos relacionados.")
p.paragraph_format.space_after = Pt(8)
p.runs[0].italic = True

security = [
"¿Tiene una enfermedad contagiosa de importancia para la salud pública, como tuberculosis (TB)?",
"¿Tiene documentación que demuestre que recibió las vacunas exigidas por la ley de Estados Unidos?",
"¿Tiene un trastorno mental o físico que represente o probablemente represente una amenaza para su seguridad o bienestar o el de otras personas?",
"¿Es o ha sido consumidor abusivo o adicto a las drogas?",
"¿Alguna vez ha sido arrestado o condenado por alguna infracción o delito, aunque haya recibido perdón, amnistía u otra medida similar?",
"¿Alguna vez ha violado o participado en una conspiración para violar alguna ley relacionada con sustancias controladas?",
"¿Es cónyuge, hijo o hija de una persona que haya violado una ley sobre tráfico de sustancias controladas y se ha beneficiado a sabiendas de esas actividades durante los últimos cinco años?",
"¿Viene a Estados Unidos para participar en prostitución o vicio comercial ilícito, o ha participado en prostitución o en la captación de prostitutas durante los últimos diez años?",
"¿Alguna vez ha participado o busca participar en lavado de dinero?",
"¿Alguna vez ha cometido o conspirado para cometer un delito de trata de personas dentro o fuera de Estados Unidos?",
"¿Alguna vez ayudó, instigó, asistió o colaboró a sabiendas con una persona identificada por el Presidente de Estados Unidos como participante importante en una forma grave de trata de personas?",
"¿Es cónyuge, hijo o hija de una persona que haya cometido o conspirado para cometer trata de personas, y durante los últimos cinco años se ha beneficiado a sabiendas de esas actividades?",
"¿Busca participar en espionaje, sabotaje, violaciones de controles de exportación u otra actividad ilegal mientras esté en Estados Unidos?",
"¿Busca participar en actividades terroristas mientras esté en Estados Unidos o ha participado alguna vez en actividades terroristas?",
"¿Alguna vez ha proporcionado o pretende proporcionar asistencia financiera u otro apoyo a terroristas u organizaciones terroristas?",
"¿Es miembro o representante de una organización terrorista?",
"¿Es cónyuge, hijo o hija de una persona que, durante los últimos cinco años, haya participado en actividades terroristas, incluido apoyo financiero o de otro tipo?",
"¿Alguna vez ordenó, incitó, cometió, ayudó o participó de otro modo en genocidio?",
"¿Alguna vez cometió, ordenó, incitó, ayudó o participó de otro modo en tortura?",
"¿Alguna vez cometió, ordenó, incitó, ayudó o participó de otro modo en ejecuciones extrajudiciales, asesinatos políticos u otros actos de violencia?",
"¿Alguna vez participó en el reclutamiento o uso de niños soldados?",
"Mientras se desempeñaba como funcionario gubernamental, ¿fue responsable o ejecutó directamente violaciones particularmente graves de la libertad religiosa?",
"¿Es miembro o está afiliado al Partido Comunista o a otro partido totalitario?",
"¿Alguna vez ayudó o apoyó, directa o indirectamente, a las FARC, el ELN o las AUC de Colombia?",
"¿Alguna vez, abusando de un cargo gubernamental o político, convirtió para beneficio personal, confiscó o expropió bienes en un país extranjero sobre los cuales un ciudadano estadounidense tenía un reclamo de propiedad?",
"¿Es cónyuge, hijo menor o agente de una persona que haya realizado esa conversión, confiscación o expropiación de bienes mediante abuso de cargo gubernamental o político?",
"¿Alguna vez participó directamente en establecer o hacer cumplir controles de población que obligaran a una mujer a abortar o a una persona a esterilizarse contra su voluntad?",
"¿Alguna vez divulgó o traficó información comercial confidencial de Estados Unidos obtenida en relación con la participación estadounidense en la Convención sobre Armas Químicas?",
"¿Es cónyuge, hijo menor o agente de una persona que haya divulgado o traficado esa información comercial confidencial?",
"¿Alguna vez intentó obtener o ayudó a otra persona a obtener una visa, entrada a Estados Unidos u otro beneficio migratorio mediante fraude, declaración falsa intencional u otro medio ilegal?",
"¿Alguna vez fue expulsado o deportado de algún país?",
"¿Alguna vez retuvo fuera de Estados Unidos la custodia de un menor ciudadano estadounidense, apartándolo de una persona a quien un tribunal estadounidense había concedido la custodia legal?",
"¿Alguna vez ayudó intencionalmente a otra persona a retener de esa forma la custodia de un menor ciudadano estadounidense fuera de Estados Unidos?",
"¿Ha votado en Estados Unidos infringiendo alguna ley o reglamento?",
"¿Alguna vez renunció a la ciudadanía estadounidense con el propósito de evitar impuestos?",
"¿Asistió a una escuela pública primaria o secundaria con estatus de estudiante F después del 30 de noviembre de 1996 sin reembolsar el costo a la escuela?",
"¿Busca entrar a Estados Unidos para realizar trabajo calificado o no calificado sin haber sido certificado por el Secretario de Trabajo?",
"¿Se graduó de una escuela de medicina extranjera y busca prestar servicios médicos en Estados Unidos sin haber aprobado el examen de la Junta Nacional de Examinadores Médicos o uno equivalente?",
"¿Es trabajador de atención médica y busca ejercer en Estados Unidos sin la certificación de la Comisión de Graduados de Escuelas de Enfermería Extranjeras o de una entidad equivalente aprobada?",
"¿Es permanentemente inelegible para obtener la ciudadanía estadounidense?",
"¿Alguna vez salió de Estados Unidos para evadir el servicio militar durante un período de guerra?",
"¿Viene a Estados Unidos para practicar la poligamia?",
"¿Es un antiguo visitante de intercambio con visa J que todavía no ha cumplido el requisito de dos años de residencia en el extranjero?",
"¿Alguna vez el Secretario de Seguridad Nacional de Estados Unidos determinó que usted presentó a sabiendas una solicitud de asilo frívola?",
"¿Es probable que se convierta en carga pública después de ser admitido en Estados Unidos?",
]
for i, text in enumerate(security, 1):
    add_q(f"{i}. {text}", True)
    if i == 2:
        add_field("Si respondió No sobre las vacunas, explique", 2)

sub("Explicaciones de respuestas afirmativas")
add_field("Número(s) de pregunta y explicación completa", 8)

section("7. Número de Seguro Social")
add_q("¿Alguna vez ha solicitado un número de Seguro Social de Estados Unidos?", True)
add_field("Si respondió Sí, indique el número si lo conoce")
add_q("¿Desea que la Administración del Seguro Social le expida un número y una tarjeta de Seguro Social?", True)
add_q("¿Autoriza que la información de este formulario sea divulgada al Departamento de Seguridad Nacional, a la Administración del Seguro Social y a otras agencias del Gobierno de Estados Unidos que la necesiten para asignarle un número de Seguro Social y emitir la tarjeta, y autoriza que la Administración del Seguro Social comparta su número con el Departamento de Seguridad Nacional?", True)

section("8. Persona que ayudó a preparar la solicitud")
add_q("¿Alguien le ayudó a completar esta solicitud?", True)
for label in ["Apellidos de la persona que ayudó", "Nombres", "Nombre de la organización (si corresponde)", "Dirección", "Apartamento/unidad", "Ciudad", "Estado/provincia", "Código postal", "País/región", "Relación con usted"]:
    add_field(label)

section("9. Revisión final del solicitante")
checks = [
"Revisé todos los nombres, fechas y números exactamente como aparecen en mis documentos.",
"Incluí todas las direcciones desde los 16 años, si corresponde.",
"Incluí todos los teléfonos, correos electrónicos y redes sociales utilizados, si corresponde.",
"Incluí todos los empleos, estudios, viajes y organizaciones requeridos.",
"Expliqué por completo cada respuesta “Sí” de seguridad y antecedentes.",
"Entiendo que la solicitud oficial se completa y presenta en el sistema CEAC.",
]
for item in checks:
    p = doc.add_paragraph("[  ]  " + item)
    p.paragraph_format.space_after = Pt(7)
add_field("Firma del solicitante")
add_field("Fecha")
add_field("Firma de la persona que ayudó (si corresponde)")
add_field("Fecha")

doc.core_properties.title = "Formulario de preparación DS-260 en español"
doc.core_properties.subject = "Preguntas para contestar manualmente"
doc.core_properties.author = "SAMSARA Services"
doc.core_properties.keywords = "DS-260, visa de inmigrante, formulario, español"
doc.save(OUT)
print(OUT)
